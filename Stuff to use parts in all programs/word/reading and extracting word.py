import docx

doc = docx.Document("new word.docx")

len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[0].runs
len(doc.paragraphs[0].runs)
doc.paragraphs[0].runs.text
doc.paragraphs[0].runs[0].text
doc.paragraphs[0].runs[0].style
def gettext(filename):
    doc= docx.Document(filename)
    Text= []
    for para in doc.paragraphs:
        Text.append(para.Text)
    return '\n'.join(Text)

print(gettext('new word.docx'))

doc.paragraphs[0].runs[0].style ='QuoteChar'
#add quotechar style
doc.paragraphs[0].runs[0].underline =True
#underlines
