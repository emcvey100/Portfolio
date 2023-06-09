import docx
doc = docx.Document()
doc.add_paragraph('Hi World')
ob1= doc.add_paragraph("another paragraph")
ob1.add_run("This is added to the second paragraph")
doc.add_heading("This is a heading with size of 18", 18)

doc.add_picture('pic.jpg', width=docx.shared.Inches(5), height=docx.shared.Cm(5)



doc.save("new word.docx")
